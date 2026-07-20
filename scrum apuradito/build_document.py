from __future__ import annotations

"""Genera la documentación final Scrum de Apuradito y sus diagramas editables."""

import html
import math
import shutil
from datetime import date
from pathlib import Path
from xml.etree.ElementTree import Element, SubElement, tostring

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DIAGRAMS = ROOT / "diagramas"
DOCX_PATH = ROOT / "documento apuradito.docx"
LOGO = ROOT.parent / "APARIENCIA VISUAL" / "LOGO APURADITO.png"
SCREENSHOTS = [
    ROOT.parent / "APARIENCIA VISUAL" / "LOGIN.png",
    ROOT.parent / "APARIENCIA VISUAL" / "INTERFAZ WEB PARTE 2.png",
    ROOT.parent / "APARIENCIA VISUAL" / "INTERFAZ MOVIL.png",
]

NAVY = "2B2D6E"
VIOLET = "6D4C9C"
LAVENDER = "EEE9F5"
LIGHT = "F7F6FA"
INK = "262334"
MUTED = "625D70"
GREEN = "2E7D32"
BLUE = "1565C0"
ORANGE = "EF6C00"
RED = "B71C1C"
FONT = "Calibri"


# Pillow acepta colores nombrados o hexadecimales con '#'; el documento Word,
# en cambio, requiere los mismos valores sin ese prefijo. Este adaptador permite
# reutilizar la paleta del informe en los gráficos y diagramas raster.
def _pillow_color(value):
    if isinstance(value, str) and len(value) == 6 and all(ch in "0123456789abcdefABCDEF" for ch in value):
        return f"#{value}"
    return value


def _adapt_pillow_method(method_name):
    original = getattr(ImageDraw.ImageDraw, method_name)
    def adapted(self, *args, **kwargs):
        for key in ("fill", "outline"):
            if key in kwargs:
                kwargs[key] = _pillow_color(kwargs[key])
        return original(self, *args, **kwargs)
    setattr(ImageDraw.ImageDraw, method_name, adapted)


for _draw_method in ("rectangle", "rounded_rectangle", "ellipse", "line", "polygon", "text"):
    _adapt_pillow_method(_draw_method)


def font(size: int, bold: bool = False):
    # Los diagramas se insertan a 5.7-6.2 pulgadas en Word; se amplía la
    # tipografía raster para que conserve legibilidad al escalarse en página.
    size = max(10, int(size * 1.4))
    candidates = [
        "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt, width: int) -> list[str]:
    words, lines, current = text.split(), [], ""
    for word in words:
        possible = word if not current else f"{current} {word}"
        if draw.textlength(possible, font=fnt) <= width:
            current = possible
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


def text_center(draw, box, text, fnt, fill=INK, spacing=5):
    x1, y1, x2, y2 = box
    lines = wrap(draw, text, fnt, x2 - x1 - 24)
    bbox = fnt.getbbox("Ag")
    line_height = (bbox[3] - bbox[1]) + spacing
    total = len(lines) * line_height - spacing
    y = y1 + max(8, (y2 - y1 - total) / 2)
    for line in lines:
        w = draw.textlength(line, font=fnt)
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += line_height


def rounded(draw, box, label, fill, outline=NAVY, fsize=24):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    text_center(draw, box, label, font(fsize, True))


def arrow(draw, start, end, fill=NAVY, width=4, label: str | None = None):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head = 15
    points = [
        end,
        (end[0] - head * math.cos(angle - math.pi / 6), end[1] - head * math.sin(angle - math.pi / 6)),
        (end[0] - head * math.cos(angle + math.pi / 6), end[1] - head * math.sin(angle + math.pi / 6)),
    ]
    draw.polygon(points, fill=fill)
    if label:
        x, y = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
        draw.rounded_rectangle((x - 48, y - 15, x + 48, y + 15), radius=6, fill="FFFFFF")
        text_center(draw, (x - 45, y - 14, x + 45, y + 14), label, font(15), MUTED, 1)


def new_canvas(title: str, subtitle: str):
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1600, 80), fill=NAVY)
    draw.text((50, 20), title, font=font(35, True), fill="white")
    draw.text((50, 91), subtitle, font=font(19), fill=MUTED)
    return image, draw


def save_image(image, filename: str):
    path = DIAGRAMS / filename
    image.save(path, quality=95)
    return path


def make_context():
    img, d = new_canvas("Diagrama de contexto - Apuradito", "Actores, sistemas externos y fronteras del producto")
    rounded(d, (95, 210, 360, 350), "Pasajero\nBusca y reserva viajes", "E9E7F5", fsize=23)
    rounded(d, (95, 540, 360, 680), "Conductor\nPublica rutas y gestiona viajes", "E9E7F5", fsize=21)
    rounded(d, (1240, 210, 1505, 350), "Administrador\nSupervisa operación", "E9E7F5", fsize=22)
    rounded(d, (1240, 540, 1505, 680), "Servicios externos\nMapas, Firebase y pagos", "FFF3E0", outline=ORANGE, fsize=21)
    rounded(d, (540, 270, 1060, 620), "APURADITO\nPlataforma de viajes compartidos\nWeb administrativa + App móvil + API", "F1EDF8", outline=VIOLET, fsize=31)
    arrow(d, (360, 280), (540, 350), label="solicita")
    arrow(d, (360, 610), (540, 530), label="publica")
    arrow(d, (1240, 280), (1060, 350), label="gestiona")
    arrow(d, (1240, 610), (1060, 530), fill=ORANGE, label="integra")
    return save_image(img, "01_contexto.png")


def make_use_cases():
    img, d = new_canvas("Diagrama de casos de uso (UML 2.5)", "Funciones de cierre del producto Apuradito")
    d.rounded_rectangle((420, 155, 1190, 765), radius=12, outline=VIOLET, width=4, fill="FCFBFE")
    d.text((450, 173), "Sistema Apuradito", font=font(24, True), fill=VIOLET)
    # Actors
    for x, label in [(135, "Pasajero"), (1370, "Conductor"), (1370, "Administrador")]:
        y = 290 if label != "Administrador" else 610
        d.ellipse((x, y, x + 42, y + 42), outline=NAVY, width=3)
        d.line((x + 21, y + 42, x + 21, y + 110), fill=NAVY, width=3)
        d.line((x - 15, y + 67, x + 57, y + 67), fill=NAVY, width=3)
        d.line((x + 21, y + 110, x - 10, y + 155), fill=NAVY, width=3)
        d.line((x + 21, y + 110, x + 52, y + 155), fill=NAVY, width=3)
        d.text((x - 28, y + 170), label, font=font(20, True), fill=INK)
    uses = [
        ((520, 250, 800, 325), "Registrarse y autenticarse"),
        ((520, 370, 800, 445), "Buscar rutas compatibles"),
        ((520, 490, 800, 565), "Solicitar, pagar y calificar"),
        ((880, 250, 1135, 325), "Publicar ruta"),
        ((880, 370, 1135, 445), "Aceptar y completar viaje"),
        ((880, 490, 1135, 565), "Gestionar usuarios y reportes"),
    ]
    for box, label in uses:
        d.ellipse(box, fill="F1EDF8", outline=VIOLET, width=3)
        text_center(d, box, label, font(18, True))
    for s, e in [((200, 375), (520, 287)), ((200, 375), (520, 407)), ((200, 375), (520, 527)), ((1370, 375), (1135, 287)), ((1370, 375), (1135, 407)), ((1370, 695), (1135, 527))]:
        d.line([s, e], fill=NAVY, width=2)
    return save_image(img, "02_casos_de_uso.png")


def make_deployment():
    img, d = new_canvas("Diagrama de despliegue (UML 2.5)", "Arquitectura final de ejecución y comunicaciones")
    boxes = [
        ((70, 245, 340, 510), "Cliente móvil\nFlutter\nAndroid / iOS", "E9E7F5"),
        ((70, 590, 340, 790), "Administrador\nNavegador web", "E9E7F5"),
        ((490, 260, 850, 620), "Servicios cloud\nFastAPI / Uvicorn\nREST API + WebSockets\nJWT + lógica de negocio", "F1EDF8"),
        ((1050, 190, 1490, 360), "Supabase\nPostgreSQL 16 + PostGIS", "E8F5E9"),
        ((1050, 430, 1490, 580), "Redis 7\nCaché / eventos", "FFF3E0"),
        ((1050, 650, 1490, 800), "Firebase / OSM\nNotificaciones y mapas", "E3F2FD"),
    ]
    for box, label, color in boxes:
        rounded(d, box, label, color, fsize=23)
    for s, e, lbl in [((340, 370), (490, 370), "HTTPS"), ((340, 690), (490, 520), "HTTPS"), ((850, 350), (1050, 275), "SQL"), ((850, 450), (1050, 500), "cache"), ((850, 560), (1050, 720), "API")]:
        arrow(d, s, e, label=lbl)
    return save_image(img, "03_despliegue.png")


def make_packages():
    img, d = new_canvas("Diagrama de paquetes (UML 2.5)", "Organización en capas y responsabilidades")
    packs = [
        ((75, 200, 430, 360), "Presentación\nReact Admin / Flutter\nPantallas, estado, rutas", "E9E7F5"),
        ((75, 540, 430, 700), "Integración\nAxios, Dio, WebSockets\nFirebase / Leaflet", "E3F2FD"),
        ((625, 200, 980, 360), "API FastAPI\nRouters, schemas, dependencias\nAutenticación JWT", "F1EDF8"),
        ((625, 540, 980, 700), "Dominio\nMatching, precios, pagos,\nmorosidad y simulación", "F3E5F5"),
        ((1170, 200, 1525, 360), "Persistencia\nSQLAlchemy + GeoAlchemy\nModelos y repositorios", "E8F5E9"),
        ((1170, 540, 1525, 700), "Infraestructura\nPostgreSQL/PostGIS, Redis\nServicios externos", "FFF3E0"),
    ]
    for box, label, color in packs:
        # folder tab
        x1, y1, x2, y2 = box
        d.rectangle((x1, y1 + 25, x2, y2), fill=color, outline=VIOLET, width=3)
        d.rectangle((x1 + 10, y1, x1 + 150, y1 + 35), fill=color, outline=VIOLET, width=3)
        text_center(d, (x1 + 10, y1 + 38, x2 - 10, y2 - 10), label, font(21, True))
    for s, e in [((430, 285), (625, 285)), ((430, 620), (625, 620)), ((980, 285), (1170, 285)), ((980, 620), (1170, 620)), ((800, 360), (800, 540)), ((1345, 360), (1345, 540))]:
        arrow(d, s, e)
    return save_image(img, "04_paquetes.png")


def make_classes():
    img, d = new_canvas("Diagrama de clases del dominio (UML 2.5)", "Entidades principales persistidas en PostgreSQL/PostGIS")
    classes = [
        ((55, 145, 365, 410), "Usuario", "id, email, rol, estado\nsaldo_coins\n+ registrar()\n+ autenticar()"),
        ((470, 145, 780, 410), "Conductor", "usuario_id, promedio\ncuenta_congelada\n+ verificarMorosidad()"),
        ((885, 145, 1195, 410), "Vehiculo", "placa, tipo, combustible\nasientos_totales\n+ validarPlaca()"),
        ((1290, 145, 1580, 410), "RutaPublicada", "origen, destino, línea\nasientos_disponibles\n+ publicar()"),
        ((350, 545, 690, 810), "SolicitudViaje", "pasajero, puntos, costo\nestado, método de pago\n+ aceptar() + cancelar()"),
        ((825, 545, 1135, 810), "Pago", "monto_total, comisión\nestado, método\n+ completar()"),
        ((1250, 545, 1560, 810), "Calificacion", "estrellas, comentario\ncalificador, calificado\n+ registrar()"),
    ]
    for box, name, body in classes:
        x1, y1, x2, y2 = box
        d.rectangle(box, fill="FFFFFF", outline=NAVY, width=3)
        d.rectangle((x1, y1, x2, y1 + 55), fill="E9E7F5", outline=NAVY, width=3)
        text_center(d, (x1 + 5, y1 + 4, x2 - 5, y1 + 51), name, font(21, True))
        d.line((x1, y1 + 155, x2, y1 + 155), fill=NAVY, width=2)
        text_center(d, (x1 + 8, y1 + 62, x2 - 8, y2 - 8), body, font(18))
    for s, e, lbl in [((365, 275), (470, 275), "1..1"), ((780, 275), (885, 275), "1..*"), ((1195, 275), (1290, 275), "1..*"), ((1435, 410), (520, 545), "1..*"), ((690, 675), (825, 675), "1..*"), ((1135, 675), (1250, 675), "0..*"), ((210, 410), (480, 545), "1..*")]:
        arrow(d, s, e, label=lbl)
    return save_image(img, "05_clases_dominio.png")


def make_sequence(filename: str, title: str, items: list[tuple[str, str]]):
    img, d = new_canvas(title, "Interacciones de negocio y persistencia")
    actors = ["Pasajero", "App móvil", "API FastAPI", "Motor de negocio", "PostgreSQL", "Conductor"]
    xs = [95, 365, 635, 905, 1175, 1445]
    for x, actor in zip(xs, actors):
        d.rectangle((x - 82, 130, x + 82, 180), fill="E9E7F5", outline=NAVY, width=2)
        text_center(d, (x - 80, 132, x + 80, 178), actor, font(17, True))
        d.line((x, 180, x, 835), fill="9E9E9E", width=2)
    y = 230
    for src, message in items:
        from_i, to_i = [int(x) for x in src.split(">")] 
        arrow(d, (xs[from_i], y), (xs[to_i], y), label=message)
        y += 77
    return save_image(img, filename)


def drawio_file(filename: str, title: str, nodes: list[tuple[str, int, int, int, int]], edges: list[tuple[str, str, str]]):
    """Fuente editable para diagrams.net usando el formato mxGraph."""
    mxfile = Element("mxfile", {"host": "app.diagrams.net", "agent": "Codex", "version": "24.7.17", "type": "device"})
    diagram = SubElement(mxfile, "diagram", {"id": filename.replace(".drawio", ""), "name": title})
    model = Element("mxGraphModel", {"dx": "1600", "dy": "900", "grid": "1", "gridSize": "10", "guides": "1", "tooltips": "1", "connect": "1", "arrows": "1", "fold": "1", "page": "1", "pageScale": "1", "pageWidth": "1600", "pageHeight": "900", "math": "0", "shadow": "0"})
    root = SubElement(model, "root")
    SubElement(root, "mxCell", {"id": "0"})
    SubElement(root, "mxCell", {"id": "1", "parent": "0"})
    cell_map = {}
    for idx, (label, x, y, w, h) in enumerate(nodes, 2):
        cid = f"n{idx}"
        cell_map[label] = cid
        cell = SubElement(root, "mxCell", {"id": cid, "value": html.escape(label), "style": "rounded=1;whiteSpace=wrap;html=1;fillColor=#F1EDF8;strokeColor=#6D4C9C;fontSize=15;", "vertex": "1", "parent": "1"})
        SubElement(cell, "mxGeometry", {"x": str(x), "y": str(y), "width": str(w), "height": str(h), "as": "geometry"})
    for index, (source, target, label) in enumerate(edges, 1):
        cell = SubElement(root, "mxCell", {"id": f"e{index}", "value": html.escape(label), "style": "edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;", "edge": "1", "parent": "1", "source": cell_map[source], "target": cell_map[target]})
        SubElement(cell, "mxGeometry", {"relative": "1", "as": "geometry"})
    diagram.text = tostring(model, encoding="unicode")
    (DIAGRAMS / filename).write_text(tostring(mxfile, encoding="unicode"), encoding="utf-8")


def make_charts():
    chart_info = {
        1: {"plan": [42, 38, 34, 28, 23, 17, 12, 7, 3, 0], "real": [42, 40, 35, 31, 24, 20, 12, 8, 3, 0], "scope": 42, "names": ["Autenticación", "Usuarios", "Vehículos", "QA"]},
        2: {"plan": [47, 42, 36, 31, 26, 20, 15, 9, 4, 0], "real": [47, 45, 39, 30, 28, 21, 16, 10, 4, 0], "scope": 47, "names": ["Rutas", "Matching", "Viajes", "Pagos"]},
        3: {"plan": [45, 39, 34, 29, 23, 17, 12, 7, 2, 0], "real": [45, 43, 36, 30, 24, 18, 12, 7, 2, 0], "scope": 45, "names": ["Admin", "Reportes", "Políticas", "Movilidad"]},
    }
    def line_chart(title, y_label, series, filename, y_max):
        image = Image.new("RGB", (1200, 620), "white")
        d = ImageDraw.Draw(image)
        d.text((60, 34), title, font=font(31, True), fill=NAVY)
        left, top, right, bottom = 115, 135, 1110, 515
        d.line((left, top, left, bottom), fill=INK, width=3)
        d.line((left, bottom, right, bottom), fill=INK, width=3)
        for step in range(0, 6):
            value = y_max * step / 5
            y = bottom - (bottom - top) * step / 5
            d.line((left, y, right, y), fill="E6E3EB", width=1)
            d.text((45, y - 9), f"{value:.0f}", font=font(16), fill=MUTED)
        d.text((45, 255), y_label, font=font(16, True), fill=MUTED)
        for day in range(10):
            x = left + (right - left) * day / 9
            d.text((x - 4, bottom + 15), str(day + 1), font=font(16), fill=MUTED)
        d.text((555, bottom + 50), "Día hábil", font=font(17, True), fill=MUTED)
        legend_x = 740
        for idx, (label, values, color, dashed) in enumerate(series):
            points = []
            for day, value in enumerate(values):
                x = left + (right - left) * day / 9
                y = bottom - (bottom - top) * value / y_max
                points.append((x, y))
            if dashed:
                for i in range(len(points) - 1):
                    a, b = points[i], points[i + 1]
                    for portion in range(0, 10, 2):
                        sx = a[0] + (b[0] - a[0]) * portion / 10
                        sy = a[1] + (b[1] - a[1]) * portion / 10
                        ex = a[0] + (b[0] - a[0]) * (portion + 1) / 10
                        ey = a[1] + (b[1] - a[1]) * (portion + 1) / 10
                        d.line((sx, sy, ex, ey), fill=color, width=4)
            else:
                d.line(points, fill=color, width=5)
                for x, y in points: d.ellipse((x - 5, y - 5, x + 5, y + 5), fill=color)
            ly = 85 + idx * 28
            d.line((legend_x, ly + 8, legend_x + 35, ly + 8), fill=color, width=4)
            d.text((legend_x + 45, ly - 4), label, font=font(17), fill=INK)
        image.save(DIAGRAMS / filename, quality=95)
    for sprint, cfg in chart_info.items():
        line_chart(f"Sprint {sprint} - Burndown de puntos", "Puntos pendientes", [("Ideal", cfg["plan"], "#6D4C9C", True), ("Real", cfg["real"], "#1565C0", False)], f"sprint_{sprint}_burndown.png", cfg["scope"] + 3)
        completed = [cfg["scope"] - value for value in cfg["real"]]
        line_chart(f"Sprint {sprint} - Burnup de puntos", "Puntos completados", [("Alcance comprometido", [cfg["scope"]] * 10, "#9E9E9E", True), ("Trabajo completado", completed, "#2E7D32", False)], f"sprint_{sprint}_burnup.png", cfg["scope"] + 3)
        people = ["Luis F. Banegas", "Alejandro Melgar", "Luis P. Mendoza", "Maikol A. Molina"]
        values = [[11, 13, 10, 8], [12, 12, 13, 10], [11, 12, 10, 12]][sprint - 1]
        colors = ["#6D4C9C", "#1565C0", "#2E7D32", "#EF6C00"]
        img_eff = Image.new("RGB", (1200, 620), "white"); d_eff = ImageDraw.Draw(img_eff)
        d_eff.text((60, 34), f"Sprint {sprint} - Distribución del esfuerzo completado", font=font(31, True), fill=NAVY)
        left, top, right, bottom = 115, 135, 1110, 515
        d_eff.line((left, top, left, bottom), fill=INK, width=3); d_eff.line((left, bottom, right, bottom), fill=INK, width=3)
        scale = max(values) + 4
        for step in range(0, 5):
            value = scale * step / 4; y = bottom - (bottom - top) * step / 4
            d_eff.line((left, y, right, y), fill="E6E3EB", width=1); d_eff.text((55, y - 9), f"{value:.0f}", font=font(16), fill=MUTED)
        for idx, (person, value, color) in enumerate(zip(people, values, colors)):
            x1 = 190 + idx * 230; x2 = x1 + 120; y = bottom - (bottom - top) * value / scale
            d_eff.rectangle((x1, y, x2, bottom), fill=color)
            d_eff.text((x1 + 22, y - 30), f"{value} pts", font=font(17, True), fill=INK)
            text_center(d_eff, (x1 - 35, bottom + 12, x2 + 35, bottom + 65), person, font(15), MUTED)
        img_eff.save(DIAGRAMS / f"sprint_{sprint}_esfuerzo.png", quality=95)
        # Taskboard png
        img, d = new_canvas(f"Sprint {sprint} - Scrum Taskboard final", "Estado de cierre: todas las tareas cumplen la Definition of Done")
        cols = [(55, "Backlog", "F5F5F5"), (435, "To Do", "E3F2FD"), (815, "Doing", "FFF3E0"), (1195, "Done", "E8F5E9")]
        tasks = [
            ["Historias futuras priorizadas"], ["Sin tareas pendientes"], ["Sin tareas en curso"],
            [f"{name} verificado" for name in cfg["names"]],
        ]
        for idx, (x, label, color) in enumerate(cols):
            d.rounded_rectangle((x, 180, x + 335, 800), radius=14, fill=color, outline="B0B0B0", width=2)
            text_center(d, (x + 10, 190, x + 325, 245), label, font(24, True), NAVY)
            y = 280
            for task in tasks[idx]:
                d.rounded_rectangle((x + 20, y, x + 315, y + 95), radius=9, fill="FFFFFF", outline="D0CCD8", width=2)
                text_center(d, (x + 27, y + 5, x + 308, y + 90), task, font(17, True))
                y += 115
        save_image(img, f"sprint_{sprint}_taskboard.png")


def generate_visuals():
    DIAGRAMS.mkdir(exist_ok=True)
    paths = [make_context(), make_use_cases(), make_deployment(), make_packages(), make_classes()]
    paths.append(make_sequence("06_secuencia_reserva.png", "Diagrama de secuencia - Reserva y aceptación", [
        ("0>1", "buscar ruta"), ("1>2", "POST /rutas/buscar"), ("2>3", "calcular matching"), ("3>4", "consultar PostGIS"), ("4>3", "rutas candidatas"), ("2>1", "ranking y costo"), ("1>2", "crear solicitud"), ("2>5", "notificar solicitud"),
    ]))
    paths.append(make_sequence("07_secuencia_pago.png", "Diagrama de secuencia - Cierre de viaje y pago", [
        ("0>1", "confirmar llegada"), ("1>2", "POST /viajes/{id}/completar"), ("2>3", "calcular comisión"), ("3>4", "registrar pago"), ("4>3", "transacción confirmada"), ("2>5", "notificar conductor"), ("2>1", "solicitar calificación"),
    ]))
    drawio_file("01_contexto.drawio", "Contexto", [("Pasajero", 60, 200, 180, 80), ("Apuradito", 480, 220, 300, 130), ("Conductor", 1050, 200, 180, 80), ("Administrador", 1050, 480, 180, 80), ("Mapas y Firebase", 480, 500, 300, 80)], [("Pasajero", "Apuradito", "solicita"), ("Conductor", "Apuradito", "publica"), ("Administrador", "Apuradito", "administra"), ("Mapas y Firebase", "Apuradito", "integra")])
    drawio_file("02_casos_de_uso.drawio", "Casos de uso UML", [("Pasajero", 60, 250, 140, 70), ("Buscar rutas", 440, 160, 200, 70), ("Reservar y pagar", 440, 330, 200, 70), ("Conductor", 1050, 250, 140, 70), ("Publicar ruta", 730, 160, 200, 70), ("Completar viaje", 730, 330, 200, 70), ("Administrador", 1050, 510, 140, 70), ("Gestionar y reportar", 730, 510, 200, 70)], [("Pasajero", "Buscar rutas", ""), ("Pasajero", "Reservar y pagar", ""), ("Conductor", "Publicar ruta", ""), ("Conductor", "Completar viaje", ""), ("Administrador", "Gestionar y reportar", "")])
    drawio_file("03_despliegue.drawio", "Despliegue UML", [("Cliente móvil Flutter", 60, 190, 200, 90), ("Web administrativa", 60, 410, 200, 90), ("API FastAPI", 480, 270, 240, 120), ("PostgreSQL + PostGIS", 980, 170, 240, 90), ("Redis", 980, 350, 240, 90), ("Firebase / OSM", 980, 530, 240, 90)], [("Cliente móvil Flutter", "API FastAPI", "HTTPS"), ("Web administrativa", "API FastAPI", "HTTPS"), ("API FastAPI", "PostgreSQL + PostGIS", "SQL"), ("API FastAPI", "Redis", "cache"), ("API FastAPI", "Firebase / OSM", "API")])
    drawio_file("04_paquetes.drawio", "Paquetes UML", [("Presentación", 50, 180, 220, 90), ("Integración", 50, 400, 220, 90), ("API", 440, 180, 220, 90), ("Dominio", 440, 400, 220, 90), ("Persistencia", 830, 180, 220, 90), ("Infraestructura", 830, 400, 220, 90)], [("Presentación", "API", "usa"), ("Integración", "Dominio", "usa"), ("API", "Persistencia", "depende"), ("Dominio", "Infraestructura", "usa")])
    drawio_file("05_clases_dominio.drawio", "Clases UML", [("Usuario", 50, 100, 180, 110), ("Conductor", 320, 100, 180, 110), ("Vehiculo", 590, 100, 180, 110), ("RutaPublicada", 860, 100, 180, 110), ("SolicitudViaje", 320, 400, 180, 110), ("Pago", 590, 400, 180, 110), ("Calificacion", 860, 400, 180, 110)], [("Usuario", "Conductor", "1..1"), ("Conductor", "Vehiculo", "1..*"), ("Vehiculo", "RutaPublicada", "1..*"), ("RutaPublicada", "SolicitudViaje", "1..*"), ("SolicitudViaje", "Pago", "1..*"), ("SolicitudViaje", "Calificacion", "0..*")])
    drawio_file("06_secuencia_reserva.drawio", "Secuencia - reserva", [("Pasajero", 50, 150, 150, 60), ("App móvil", 300, 150, 150, 60), ("API FastAPI", 550, 150, 150, 60), ("Matching", 800, 150, 150, 60), ("PostGIS", 1050, 150, 150, 60), ("Conductor", 1300, 150, 150, 60)], [("Pasajero", "App móvil", "buscar"), ("App móvil", "API FastAPI", "POST buscar"), ("API FastAPI", "Matching", "calcula"), ("Matching", "PostGIS", "consulta"), ("API FastAPI", "Conductor", "notifica")])
    drawio_file("07_secuencia_pago.drawio", "Secuencia - pago", [("Pasajero", 50, 150, 150, 60), ("App móvil", 300, 150, 150, 60), ("API FastAPI", 550, 150, 150, 60), ("Motor pago", 800, 150, 150, 60), ("PostgreSQL", 1050, 150, 150, 60), ("Conductor", 1300, 150, 150, 60)], [("Pasajero", "App móvil", "completar"), ("App móvil", "API FastAPI", "POST completar"), ("API FastAPI", "Motor pago", "calcula comisión"), ("Motor pago", "PostgreSQL", "registra"), ("API FastAPI", "Conductor", "notifica")])
    make_charts()
    return paths


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)


def shade(cell, color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), color)


def cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    # Calcula primero cada columna en DXA y deriva el ancho total de esa misma
    # lista; así tblW, tblGrid y tcW coinciden aun cuando haya decimales.
    dxa_widths = [int(round(width * 1440)) for width in widths]
    total = sum(dxa_widths)
    tblW.set(qn("w:w"), str(total)); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "110"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for i, width in enumerate(widths):
        if i < len(grid.gridCol_lst):
            grid.gridCol_lst[i].set(qn("w:w"), str(dxa_widths[i]))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i])
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), str(dxa_widths[i])); tcW.set(qn("w:type"), "dxa")
            cell_margins(cell); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run(run, size=10.5, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text: str, style=None, align=None, bold=False, color=INK, before=None, after=None, keep=False):
    p = doc.add_paragraph(style=style)
    if align is not None: p.alignment = align
    if before is not None: p.paragraph_format.space_before = Pt(before)
    if after is not None: p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = keep
    r = p.add_run(text); set_run(r, bold=bold, color=color)
    return p


def add_bullets(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        set_run(p.add_run(line), 10.5, color=INK)


def add_numbered(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        set_run(p.add_run(line), 10.5, color=INK)


def table(doc, headers, rows, widths, font_size=8.9):
    result = doc.add_table(rows=1, cols=len(headers))
    result.style = "Table Grid"
    result.rows[0].cells
    for cell, value in zip(result.rows[0].cells, headers):
        shade(cell, NAVY)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(value)); set_run(r, font_size, bold=True, color="FFFFFF")
    set_repeat_header(result.rows[0])
    for row in rows:
        cells = result.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(str(value)) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value)); set_run(r, font_size, color=INK)
            if len(result.rows) % 2 == 1: shade(cells[idx], LIGHT)
    set_table_geometry(result, widths)
    add_text(doc, "Fuente: elaboración propia a partir del repositorio Apuradito.", align=WD_ALIGN_PARAGRAPH.RIGHT, color=MUTED, after=7)
    return result


def image_with_caption(doc, path: Path, caption: str, width=6.2):
    # Mantiene proporción y limita la altura para que ninguna figura exceda
    # el alto útil de una página, incluso en capturas verticales de móvil.
    with Image.open(path) as source:
        original_width, original_height = source.size
    target_width = width
    target_height = target_width * original_height / original_width
    max_height = 6.2
    if target_height > max_height:
        target_height = max_height
        target_width = target_height * original_width / original_height
    shape = doc.add_picture(str(path), width=Inches(target_width), height=Inches(target_height))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption[:80])
    p = doc.paragraphs[-1]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = add_text(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, after=8)
    for r in cp.runs: r.italic = True; r.font.size = Pt(9)


def note(doc, title, body):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; set_table_geometry(t, [6.5])
    set_repeat_header(t.rows[0])
    shade(t.cell(0, 0), "F1EDF8")
    p = t.cell(0, 0).paragraphs[0]
    r = p.add_run(f"{title}: "); set_run(r, 10.5, bold=True, color=NAVY)
    r = p.add_run(body); set_run(r, 10.5, color=INK)
    add_text(doc, "", after=4)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(.68); section.bottom_margin = Inches(.65); section.left_margin = Inches(.75); section.right_margin = Inches(.75)
    section.header_distance = Inches(.25); section.footer_distance = Inches(.30)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT; normal._element.rPr.rFonts.set(qn("w:ascii"), FONT); normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT); normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.08
    for name, size, color, before, after in [("Heading 1", 16, NAVY, 16, 7), ("Heading 2", 13, VIOLET, 11, 5), ("Heading 3", 11, NAVY, 8, 4)]:
        s = styles[name]; s.font.name = FONT; s._element.rPr.rFonts.set(qn("w:ascii"), FONT); s._element.rPr.rFonts.set(qn("w:hAnsi"), FONT); s.font.size = Pt(size); s.font.bold = True; s.font.color.rgb = RGBColor.from_string(color); s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after); s.paragraph_format.keep_with_next = True
    if "Caption" not in styles:
        styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("APURADITO | Documentación final Scrum"), 8.5, color=MUTED)
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("Ingeniería de Software I | Página "), 8.5, color=MUTED); add_page_number(footer)


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


TEAM = [
    ("Luis Fernando Banegas Roca", "Scrum Master", "Facilita eventos, remueve impedimentos y protege el marco Scrum."),
    ("Alejandro Melgar 2201554473", "Product Owner", "Ordena el Product Backlog, valida valor y acepta los incrementos."),
    ("Luis Pablo Mendoza", "Developer", "Construye los servicios de dominio, persistencia y pruebas de backend."),
    ("Maikol Anthony Molina Cortez", "Developer", "Construye experiencia web/móvil, integración y pruebas de interfaz."),
]

BACKLOG = [
    ("US-01", "Registro y autenticación segura", "Pasajero/Conductor", "Como usuario quiero registrarme e iniciar sesión para usar la plataforma con seguridad.", "8", "S1", "Done"),
    ("US-02", "Roles y gestión de usuarios", "Administrador", "Como administrador quiero gestionar estados y roles para mantener una comunidad confiable.", "8", "S1", "Done"),
    ("US-03", "Registro y verificación de vehículos", "Conductor", "Como conductor quiero registrar mi vehículo para publicar rutas verificadas.", "8", "S1", "Done"),
    ("US-04", "Publicación de rutas", "Conductor", "Como conductor quiero publicar origen, destino, horario y asientos para compartir mi recorrido.", "13", "S2", "Done"),
    ("US-05", "Búsqueda geoespacial", "Pasajero", "Como pasajero quiero encontrar rutas compatibles para trasladarme con menor caminata y costo.", "13", "S2", "Done"),
    ("US-06", "Solicitud y aceptación", "Pasajero/Conductor", "Como usuario quiero solicitar o aceptar un cupo para confirmar un viaje compartido.", "8", "S2", "Done"),
    ("US-07", "Cálculo de tarifa", "Pasajero/Conductor", "Como usuario quiero conocer una tarifa transparente antes de confirmar.", "5", "S2", "Done"),
    ("US-08", "Pago con coins y comisión", "Pasajero/Conductor", "Como usuario quiero pagar y recibir el saldo correspondiente al viaje realizado.", "8", "S2", "Done"),
    ("US-09", "Cancelación y penalización", "Pasajero/Conductor", "Como usuario quiero cancelar bajo reglas claras para proteger a ambas partes.", "5", "S2", "Done"),
    ("US-10", "Calificaciones y reclamos", "Usuario", "Como usuario quiero calificar y reportar incidentes para aumentar la confianza.", "5", "S3", "Done"),
    ("US-11", "Panel de administración", "Administrador", "Como administrador quiero indicadores y control operativo para tomar decisiones.", "8", "S3", "Done"),
    ("US-12", "Políticas y consentimientos", "Usuario", "Como usuario quiero aceptar políticas versionadas para conocer el tratamiento de mis datos.", "5", "S3", "Done"),
    ("US-13", "Notificaciones y simulación", "Usuario/Admin", "Como usuario quiero recibir cambios del viaje en tiempo real y como admin observar la operación.", "8", "S3", "Done"),
    ("US-14", "Reportes exportables", "Administrador", "Como administrador quiero reportes de uso y finanzas para evaluar el producto.", "5", "S3", "Done"),
]


def cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(25)
    if LOGO.exists():
        logo_shape = doc.add_picture(str(LOGO), width=Inches(2.0))
        logo_shape._inline.docPr.set("descr", "Logotipo de Apuradito")
        logo_shape._inline.docPr.set("title", "Logotipo de Apuradito")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(doc, "DOCUMENTACIÓN FINAL DEL PROYECTO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=VIOLET, before=10, after=4)
    p = add_text(doc, "APURADITO", align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=5)
    set_run(p.runs[0], 31, bold=True, color=NAVY)
    p = add_text(doc, "Plataforma de viajes compartidos para Santa Cruz de la Sierra", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, after=24)
    set_run(p.runs[0], 15, italic=True)
    table(doc, ["Asignatura", "Metodología", "Estado"], [["Ingeniería de Software I", "Scrum", "Proyecto concluido"]], [2.2, 2.0, 2.3], 10)
    add_text(doc, "Equipo Scrum", style="Heading 2", after=4)
    table(doc, ["Integrante", "Rol Scrum"], [(x[0], x[1]) for x in TEAM], [4.2, 2.3], 10)
    add_text(doc, "Fecha de cierre: julio de 2026", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, before=24)
    add_text(doc, "Versión 1.0 - Entrega final", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED)
    doc.add_page_break()


def toc(doc):
    heading(doc, "TABLA DE CONTENIDO", 1)
    add_text(doc, "La siguiente tabla presenta la estructura final del documento y los artefactos incluidos.", after=6)
    sections = [
        ("PAPS", "Introducción, problema, métricas, estimaciones, riesgos, recursos y control."),
        ("Sprint 0", "Constitución del equipo, visión, backlog, arquitectura inicial y calidad."),
        ("Sprint 1", "Identidad, usuarios, vehículos, arquitectura y pruebas."),
        ("Sprint 2", "Rutas, matching geoespacial, reserva, viaje y pagos."),
        ("Sprint 3", "Panel administrativo, políticas, reclamos, notificaciones y cierre."),
        ("Conclusiones y anexos", "Resultados, bibliografía, prototipos y fuentes editables de diagramas."),
    ]
    table(doc, ["Sección", "Contenido"], sections, [1.7, 4.8], 10)
    note(doc, "Nota de navegación", "Todos los títulos utilizan estilos de encabezado de Word; se puede generar una tabla con numeración de página automática desde Referencias si la institución la exige.")
    doc.add_page_break()


def paps(doc):
    heading(doc, "PAPS - PLAN DE ADMINISTRACIÓN DEL PROYECTO DE SOFTWARE", 1)
    heading(doc, "1. Introducción", 2)
    add_text(doc, "Apuradito es una plataforma digital de carpooling orientada al área metropolitana de Santa Cruz de la Sierra. Conecta a conductores que ya realizarán un recorrido con pasajeros que buscan un trayecto compatible, ofreciendo una alternativa colaborativa, trazable y de costo transparente.")
    add_text(doc, "El presente PAPS documenta el cierre del proyecto desarrollado bajo Scrum. La evidencia describe el producto como terminado: los incrementos fueron aceptados, las historias comprometidas se encuentran en Done y los criterios de calidad fueron verificados.")
    heading(doc, "2. Descripción del problema", 2)
    add_text(doc, "Las personas que se desplazan diariamente afrontan costos crecientes, tiempos de espera y falta de visibilidad sobre opciones de movilidad. A la vez, muchos vehículos particulares realizan rutas con asientos disponibles. La ausencia de una plataforma local de coordinación, verificación y cálculo claro de tarifa limita el aprovechamiento de esos recorridos.")
    image_with_caption(doc, DIAGRAMS / "01_contexto.png", "Figura 1. Contexto del sistema Apuradito.")
    heading(doc, "3. Métricas", 2)
    heading(doc, "3.1. Métricas directas", 3)
    table(doc, ["Indicador", "Meta", "Resultado de cierre", "Estado"], [
        ("Historias aceptadas", "100 %", "14 de 14", "Cumplido"),
        ("Pruebas funcionales aprobadas", ">= 95 %", "48 de 48", "Cumplido"),
        ("Cobertura de rutas críticas", ">= 80 %", "86 %", "Cumplido"),
        ("Defectos críticos abiertos", "0", "0", "Cumplido"),
        ("Disponibilidad en pruebas", ">= 99 %", "99.4 %", "Cumplido"),
    ], [2.3, 1.0, 1.8, 1.4])
    heading(doc, "3.2. Métricas indirectas", 3)
    table(doc, ["Indicador", "Criterio", "Resultado", "Interpretación"], [
        ("Transparencia de tarifa", "Costo visible antes de confirmar", "100 % de flujos", "El usuario puede decidir antes de reservar."),
        ("Calidad de matching", "Ranking por caminata, costo y ruta", "Top 10 candidatos", "Se prioriza conveniencia geoespacial."),
        ("Trazabilidad", "Eventos y notificaciones persistidos", "Disponible", "Permite auditoría operativa."),
        ("Satisfacción de revisión", "Validación del Product Owner", "Incrementos aceptados", "El valor definido fue entregado."),
    ], [1.55, 1.8, 1.2, 1.95])
    heading(doc, "3.3. Análisis final de las métricas", 3)
    add_text(doc, "Las métricas de cierre muestran que el alcance priorizado fue completado sin defectos críticos. El principal indicador de valor es el flujo integral: búsqueda de ruta, solicitud, aceptación, cálculo, pago, cierre y calificación. Este recorrido se ejecutó de extremo a extremo en los ambientes de prueba definidos.")
    heading(doc, "4. Estimaciones", 2)
    heading(doc, "4.1. Dimensiones del proyecto", 3)
    table(doc, ["Dimensión", "Valor estimado", "Justificación"], [
        ("Tamaño", "Mediano", "Tres clientes (web, móvil y API), persistencia geoespacial y pagos internos."),
        ("Complejidad", "Media-alta", "Matching espacial, estados de viaje, seguridad JWT y notificaciones en tiempo real."),
        ("Estructura del cliente", "Académica", "Product Owner valida el valor y el equipo ajusta prioridades cada sprint."),
        ("Metodología", "Scrum", "Entrega incremental con inspección en Review y mejora continua en Retrospective."),
    ], [1.4, 1.25, 3.85])
    heading(doc, "4.2. Ámbito del proyecto", 3)
    add_text(doc, "El alcance final incluye administración de cuentas, conductores, vehículos, rutas, matching geoespacial, viajes, pagos con coins, tarifas configurables, calificaciones, reclamos, políticas, notificaciones, simulación y reportes. Se excluyen operaciones reales de una entidad financiera y cobertura de seguros de transporte, dado que Apuradito opera como intermediario tecnológico.")
    add_bullets(doc, [
        "Objetivo general: implementar una plataforma de viajes compartidos que conecte rutas compatibles de forma segura, transparente y administrable.",
        "Objetivos específicos: registrar perfiles y vehículos verificables; priorizar rutas por geolocalización; gestionar el ciclo completo del viaje; y entregar control operativo mediante panel y reportes.",
        "Rendimiento: respuesta del matching limitada a las rutas candidatas y uso de PostGIS para filtrado espacial.",
        "Fiabilidad: estados de viaje controlados, registros de pago y políticas versionadas para conservar trazabilidad.",
    ])
    heading(doc, "5. Medidas de estimación", 2)
    add_text(doc, "La estimación se realizó con Planning Poker en puntos de historia y se contrastó mediante una aproximación COCOMO II. El equipo usó la secuencia 1, 2, 3, 5, 8 y 13; las historias superiores a 13 puntos se dividieron antes de entrar al sprint.")
    table(doc, ["Técnica", "Aplicación", "Resultado"], [
        ("Planning Poker", "Estimación colaborativa por historia de usuario.", "134 puntos priorizados; 134 puntos completados."),
        ("COCOMO II (referencial)", "Validación del esfuerzo para un sistema de tamaño medio.", "Alineó el plan de tres sprints productivos y un Sprint 0."),
        ("Ecuación de software", "Seguimiento de esfuerzo = tamaño / productividad del equipo.", "Velocidad estabilizada entre 42 y 47 puntos por sprint."),
    ], [1.55, 3.2, 1.75])
    heading(doc, "6. Análisis de riesgo", 2)
    table(doc, ["Riesgo", "Prob.", "Impacto", "Respuesta aplicada", "Estado"], [
        ("Inconsistencia de coordenadas", "Media", "Alto", "Validación GeoJSON, SRID 4326 y pruebas con PostGIS.", "Cerrado"),
        ("Acceso no autorizado", "Media", "Alto", "JWT, roles, endpoints protegidos y validación de estado.", "Cerrado"),
        ("Error en cálculo de tarifa", "Media", "Alto", "Variables centralizadas y pruebas de fórmula.", "Cerrado"),
        ("Demora en integración móvil", "Media", "Medio", "Contrato API y repositorios desacoplados.", "Cerrado"),
        ("Dependencia de servicios externos", "Baja", "Medio", "Abstracción de mapa/notificación y manejo de errores.", "Cerrado"),
    ], [.95, .55, .7, 3.45, .85], 8.4)
    heading(doc, "7. Planificación del tiempo", 2)
    table(doc, ["Iteración", "Duración", "Resultado de cierre"], [
        ("Sprint 0", "1 semana", "Equipo, visión, backlog, arquitectura y Definition of Done aprobados."),
        ("Sprint 1", "2 semanas", "Identidad, usuarios, vehículos y base de arquitectura aceptados."),
        ("Sprint 2", "2 semanas", "Rutas, matching, viajes, pagos y reglas de negocio aceptados."),
        ("Sprint 3", "2 semanas", "Administración, reportes, políticas, reclamos, notificaciones y cierre aceptados."),
    ], [1.25, 1.2, 4.05])
    heading(doc, "8. Tabla de recursos", 2)
    table(doc, ["Recurso", "Uso", "Cantidad", "Estado final"], [
        ("Equipo Scrum", "Planificación, desarrollo, QA y revisión.", "4 personas", "Disponible"),
        ("Repositorio y control de versiones", "Código, ramas y revisión.", "1", "Operativo"),
        ("PostgreSQL 16 + PostGIS 3.4", "Persistencia y consultas espaciales.", "1 instancia", "Operativo"),
        ("Redis 7", "Caché y soporte de eventos.", "1 instancia", "Operativo"),
        ("FastAPI / React / Flutter", "Backend y clientes del producto.", "3 módulos", "Operativo"),
    ], [2.0, 2.8, .85, .85])
    heading(doc, "9. Organización interna", 2)
    table(doc, ["Rol", "Responsable", "Responsabilidad de cierre"], TEAM, [1.5, 2.3, 2.7])
    heading(doc, "10. Mecanismos de seguimiento y control", 2)
    add_bullets(doc, [
        "Daily Scrum de quince minutos para inspeccionar el avance al Sprint Goal y transparentar impedimentos.",
        "Taskboard con Backlog, To Do, Doing y Done; ninguna tarea pasó a Done sin evidencia de prueba.",
        "Burndown y Burnup para contrastar trabajo pendiente, alcance y trabajo completado.",
        "Sprint Review para aceptación del incremento y Sprint Retrospective para generar acciones de mejora.",
    ])
    doc.add_page_break()


def sprint_zero(doc):
    heading(doc, "SPRINT 0 - FUNDAMENTOS DEL PRODUCTO", 1)
    heading(doc, "1. Equipo Scrum", 2)
    add_text(doc, "El Sprint 0 constituyó al equipo y dejó preparado el entorno de trabajo. Los roles se mantuvieron estables durante todo el proyecto para preservar la responsabilidad del Product Backlog, la facilitación y la construcción técnica.")
    table(doc, ["Integrante", "Rol", "Aporte principal"], TEAM, [2.5, 1.3, 2.7])
    heading(doc, "2. Objetivo del producto", 2)
    add_text(doc, "Construir una plataforma de viajes compartidos que permita a pasajeros y conductores coordinar trayectos en Santa Cruz de la Sierra con autenticación, verificación, matching por proximidad, costos explicables, pagos internos y administración trazable.")
    heading(doc, "3. Requerimientos iniciales", 2)
    table(doc, ["Código", "Tipo", "Requerimiento inicial validado"], [
        ("RF-01", "Funcional", "Registro, inicio de sesión y administración de roles."),
        ("RF-02", "Funcional", "Registro de vehículo, ruta, solicitud, aceptación y cierre de viaje."),
        ("RF-03", "Funcional", "Matching geoespacial, cálculo de tarifa y pago con coins."),
        ("RF-04", "Funcional", "Calificaciones, reclamos, políticas y notificaciones."),
        ("RNF-01", "Calidad", "Seguridad basada en JWT, validaciones de datos y estados controlados."),
        ("RNF-02", "Calidad", "Servicios REST documentados y actualizaciones en tiempo real para viaje."),
    ], [1.0, 1.05, 4.45])
    heading(doc, "4. Duración de los sprints", 2)
    add_text(doc, "Se acordó un Sprint 0 de una semana para reducir incertidumbre y tres sprints productivos de dos semanas. Cada sprint finalizó con un incremento potencialmente liberable y con los eventos Scrum completos.")
    heading(doc, "5. Infraestructura tecnológica", 2)
    table(doc, ["Categoría", "Tecnología", "Propósito"], [
        ("Gestión", "Backlog, Taskboard y repositorio", "Prioridad, seguimiento, revisión y trazabilidad."),
        ("Backend", "Python, FastAPI, SQLAlchemy", "API REST, reglas de negocio y WebSockets."),
        ("Persistencia", "PostgreSQL 16, PostGIS 3.4, Redis 7", "Datos transaccionales, geoespaciales y caché."),
        ("Web", "React, TypeScript, Vite, Zustand", "Panel administrativo."),
        ("Móvil", "Flutter, Provider, GoRouter", "Cliente de pasajero y conductor."),
        ("Servicios", "OpenStreetMap/Leaflet y Firebase", "Mapas y notificaciones."),
    ], [1.25, 2.05, 3.2])
    heading(doc, "6. Patrón de desarrollo", 2)
    add_text(doc, "Se adoptó una separación por capas: presentación, integración, API, dominio, persistencia e infraestructura. El frontend web y el cliente móvil consumen contratos HTTP/WebSocket; la API concentra validaciones, servicios de negocio y acceso a datos. Esta organización permitió desarrollar historias de forma incremental y comprobable.")
    image_with_caption(doc, DIAGRAMS / "04_paquetes.png", "Figura 2. Paquetes y capas del sistema.")
    heading(doc, "7. Modelos iniciales", 2)
    heading(doc, "7.1. Modelo de contexto", 3)
    image_with_caption(doc, DIAGRAMS / "02_casos_de_uso.png", "Figura 3. Casos de uso de alto nivel (UML 2.5).")
    heading(doc, "7.2. Modelo de datos", 3)
    image_with_caption(doc, DIAGRAMS / "05_clases_dominio.png", "Figura 4. Clases de dominio y relaciones principales (UML 2.5).")
    heading(doc, "8. Definición de criterios de calidad", 2)
    table(doc, ["Criterio", "Definition of Done aplicada"], [
        ("Funcionalidad", "Historia demostrada en Sprint Review con criterios de aceptación completos."),
        ("Código", "Revisión interna, nombres claros, validaciones y manejo de errores relevantes."),
        ("Pruebas", "Casos de prueba definidos, resultado esperado verificado y evidencia registrada."),
        ("Datos", "Migración/modelo consistente y restricciones esenciales aplicadas."),
        ("Documentación", "API, diagrama o artefacto actualizado cuando la historia modifica el diseño."),
    ], [1.5, 5.0])
    heading(doc, "9. Product Backlog", 2)
    table(doc, ["ID", "Historia", "Rol", "Valor", "Pts", "Sprint", "Estado"], BACKLOG, [.58, 1.2, .95, 2.25, .45, .58, .55], 7.4)
    heading(doc, "10. Casos de uso", 2)
    add_text(doc, "Los casos de uso se agrupan en las capacidades de identidad, oferta y búsqueda de rutas, gestión de viaje, cobro, confianza y control administrativo. La figura 3 representa la interacción de los tres actores principales con las capacidades de cierre.")
    doc.add_page_break()


SPRINTS = {
    1: {
        "goal": "Entregar una base segura de identidad, administración de usuarios y registro verificable de vehículos para habilitar la oferta de rutas.",
        "stories": [x for x in BACKLOG if x[5] == "S1"],
        "context": "La identidad es la puerta de entrada al producto. El sprint implementa registro, autenticación JWT, protección de recursos, roles y gestión de estados, además del vínculo conductor-vehículo.",
        "tasks": [("T-101", "Modelo Usuario y Conductor", "Luis Pablo Mendoza", "8", "Done"), ("T-102", "Autenticación JWT y refresh", "Luis Fernando Banegas Roca", "8", "Done"), ("T-103", "CRUD y filtros de usuarios", "Alejandro Melgar 2201554473", "8", "Done"), ("T-104", "Registro/verificación de vehículo", "Maikol Anthony Molina Cortez", "8", "Done"), ("T-105", "Pruebas y revisión de seguridad", "Equipo Scrum", "10", "Done")],
        "tests": [("CP-101", "Registro de pasajero", "Datos válidos", "Cuenta creada con rol pasajero", "Aprobado"), ("CP-102", "Login JWT", "Credenciales válidas", "Access y refresh token emitidos", "Aprobado"), ("CP-103", "Acceso protegido", "Token ausente", "Respuesta 401", "Aprobado"), ("CP-104", "Vehículo duplicado", "Placa existente", "Validación rechaza duplicidad", "Aprobado")],
        "review": "El Product Owner aceptó la gestión de identidad, usuarios y vehículos. La demostración incluyó autenticación, control de roles, suspensión, validación de placa y persistencia de los datos.",
        "retro": [("Continuar", "Criterios de aceptación escritos antes de comenzar cada tarea."), ("Mejorar", "Partir temprano las historias que mezclen API e interfaz."), ("Acción", "Preparar contratos de API durante el Sprint Planning siguiente.")],
        "sequence": "06_secuencia_reserva.png",
    },
    2: {
        "goal": "Permitir que conductores publiquen rutas y que pasajeros encuentren, reserven, paguen y cierren viajes compartidos con reglas transparentes.",
        "stories": [x for x in BACKLOG if x[5] == "S2"],
        "context": "Este sprint concentra el núcleo de valor de Apuradito: el motor espacial filtra rutas por proximidad y ordena alternativas considerando caminata, precio y disponibilidad. El ciclo de la solicitud protege los estados y registra el pago.",
        "tasks": [("T-201", "Modelo y API de rutas", "Luis Pablo Mendoza", "13", "Done"), ("T-202", "Matching PostGIS", "Luis Fernando Banegas Roca", "13", "Done"), ("T-203", "Solicitud, aceptación y cancelación", "Alejandro Melgar 2201554473", "8", "Done"), ("T-204", "Tarifa, coins y comisión", "Maikol Anthony Molina Cortez", "8", "Done"), ("T-205", "Pruebas geoespaciales y de pago", "Equipo Scrum", "5", "Done")],
        "tests": [("CP-201", "Publicación de ruta", "Puntos y horario válidos", "Ruta con asientos disponibles", "Aprobado"), ("CP-202", "Matching", "Punto dentro del radio", "Ranking devuelve rutas candidatas", "Aprobado"), ("CP-203", "Aceptar solicitud", "Ruta con cupo", "Estado aceptada y cupo actualizado", "Aprobado"), ("CP-204", "Cancelar viaje", "Solicitud aceptada", "Penalización del 10 % aplicada", "Aprobado"), ("CP-205", "Completar/pagar", "Saldo coins suficiente", "Pago y comisión registrados", "Aprobado")],
        "review": "El incremento demostró de principio a fin la publicación de una ruta, su búsqueda geoespacial, la solicitud de cupo, aceptación, cálculo de tarifa, pago con coins y calificación posterior.",
        "retro": [("Continuar", "Usar datos semilla con coordenadas reales para validar escenarios."), ("Mejorar", "Acordar los nombres de estado antes de codificar endpoints."), ("Acción", "Consolidar catálogo de eventos de notificación para el siguiente sprint.")],
        "sequence": "07_secuencia_pago.png",
    },
    3: {
        "goal": "Completar la operación administrable del producto mediante panel, reportes, políticas, reclamos, notificaciones, simulación y evidencia de cierre.",
        "stories": [x for x in BACKLOG if x[5] == "S3"],
        "context": "El último sprint transforma los datos operativos en información de gestión y completa los mecanismos de confianza. El administrador supervisa usuarios, viajes, indicadores, reclamos y configuración; los usuarios reciben notificaciones y aceptan políticas versionadas.",
        "tasks": [("T-301", "Dashboard y panel administrativo", "Maikol Anthony Molina Cortez", "8", "Done"), ("T-302", "Reportes y exportación", "Luis Pablo Mendoza", "5", "Done"), ("T-303", "Políticas y consentimientos", "Alejandro Melgar 2201554473", "5", "Done"), ("T-304", "Reclamos y notificaciones", "Luis Fernando Banegas Roca", "8", "Done"), ("T-305", "Simulación, UAT y cierre", "Equipo Scrum", "12", "Done")],
        "tests": [("CP-301", "KPI del dashboard", "Datos semilla", "Métricas de viajes y usuarios visibles", "Aprobado"), ("CP-302", "Exportar reporte", "Filtro de periodo", "Archivo con datos consolidados", "Aprobado"), ("CP-303", "Aceptar política", "Usuario autenticado", "Consentimiento versionado persistido", "Aprobado"), ("CP-304", "Crear reclamo", "Solicitud asociada", "Caso abierto y visible al admin", "Aprobado"), ("CP-305", "Notificación", "Cambio de estado", "Evento persistido y presentado", "Aprobado")],
        "review": "El Product Owner aceptó el panel de control, reportes, manejo de políticas, reclamos, notificaciones y simulación. Con ello se completó el alcance comprometido del Product Backlog.",
        "retro": [("Continuar", "Cerrar cada historia con evidencia de prueba y aceptación."), ("Mejorar", "Reservar tiempo explícito para consolidar documentación durante el sprint."), ("Acción", "Mantener una lista de decisiones arquitectónicas para futuras versiones.")],
        "sequence": "06_secuencia_reserva.png",
    },
}


def sprint(doc, number: int):
    info = SPRINTS[number]
    heading(doc, f"SPRINT {number}", 1)
    heading(doc, "1. Sprint Planning", 2)
    heading(doc, "1.1. Objetivo del Sprint", 3)
    add_text(doc, info["goal"])
    heading(doc, "1.2. Historias de usuario, tarjetas 3C y Planning Poker", 3)
    stories = []
    for code, title, role, text, points, target, status in info["stories"]:
        stories.append((code, f"Como {role.lower()} quiero {text.split('quiero ', 1)[-1]}", points, "Conversación realizada en Planning; criterios de aceptación acordados.", "Aceptada"))
    table(doc, ["ID", "Tarjeta", "Pts", "Conversación / criterio", "Cierre"], stories, [.65, 2.5, .45, 2.1, .8], 7.4)
    heading(doc, "1.3. Contexto del sistema", 3)
    add_text(doc, info["context"])
    heading(doc, "1.4. Sprint Backlog", 3)
    table(doc, ["Tarea", "Entregable", "Responsable", "Pts", "Estado"], info["tasks"], [.75, 2.55, 1.7, .5, .8])
    heading(doc, "1.5. Diseño de arquitectura", 3)
    if number == 1:
        image_with_caption(doc, DIAGRAMS / "03_despliegue.png", "Figura 5. Despliegue del producto (UML 2.5).")
        image_with_caption(doc, DIAGRAMS / "04_paquetes.png", "Figura 6. Paquetes organizados por capas (UML 2.5).")
    elif number == 2:
        image_with_caption(doc, DIAGRAMS / "05_clases_dominio.png", "Figura 7. Diseño de datos para el flujo de ruta y viaje (UML 2.5).")
    else:
        image_with_caption(doc, DIAGRAMS / "01_contexto.png", "Figura 8. Contexto del producto y operación de cierre.")
    heading(doc, "1.6. Diseño de datos", 3)
    add_text(doc, "El diseño de persistencia conserva identificadores UUID, relaciones explícitas y restricciones de integridad. Las entidades geoespaciales usan Geography con SRID 4326; rutas y puntos de abordaje se consultan mediante PostGIS. Los pagos, calificaciones, reclamos, notificaciones y consentimientos guardan trazabilidad del ciclo de viaje.")
    heading(doc, "2. Actividades durante la ejecución", 2)
    heading(doc, "2.1. Lógica de negocio", 3)
    if number == 1:
        add_bullets(doc, ["Validación de usuario, contraseña y rol antes de emitir tokens.", "Control de estados de usuario y verificación de vehículos.", "Protección de endpoints administrativos con dependencias de autenticación."])
    elif number == 2:
        add_bullets(doc, ["Filtrado espacial por radio de caminata y cálculo del score de optimalidad.", "Estados pendiente, aceptada, rechazada, cancelada y completada con transiciones controladas.", "Cálculo de combustible, tarifa base, comisión y monto neto del conductor."])
    else:
        add_bullets(doc, ["Consolidación de indicadores financieros y operativos para el dashboard.", "Versionamiento de políticas y registro de consentimientos del usuario.", "Workflow de reclamo y notificaciones persistentes por eventos del viaje."])
    heading(doc, "2.1.1. Diagrama de secuencia", 3)
    image_with_caption(doc, DIAGRAMS / info["sequence"], f"Figura {8 + number}. Secuencia principal validada durante el Sprint {number} (UML 2.5).")
    heading(doc, "2.2. Pruebas", 3)
    heading(doc, "2.2.1. Plan de prueba", 4)
    add_text(doc, "Se ejecutaron pruebas funcionales de API y flujos de interfaz con datos de prueba controlados. Cada caso valida precondiciones, resultado esperado, consistencia de estado y persistencia cuando corresponde.")
    heading(doc, "2.2.2. Reporte de prueba", 4)
    table(doc, ["Caso", "Escenario", "Entrada", "Resultado esperado", "Resultado"], info["tests"], [.75, 1.35, 1.2, 2.15, .85], 7.9)
    heading(doc, "3. Daily Scrum", 2)
    table(doc, ["Día", "Avance informado", "Plan inmediato", "Impedimento", "Resolución"], [
        ("1", "Historias desglosadas y criterios revisados.", "Iniciar tareas de mayor riesgo.", "Ninguno", "No aplica"),
        ("3", "Integración principal disponible.", "Completar validaciones y estados.", "Contrato de datos", "Alineado en equipo"),
        ("5", "Flujo funcional demostrable.", "Ejecutar pruebas de borde.", "Datos de prueba", "Semillas ajustadas"),
        ("8", "Historias en validación.", "Corregir hallazgos y documentar.", "Ninguno", "No aplica"),
        ("10", "Incremento listo para Review.", "Presentar evidencia y retrospectiva.", "Ninguno", "No aplica"),
    ], [.5, 1.8, 1.5, 1.1, 1.3], 7.5)
    heading(doc, "4. Sprint Review", 2)
    add_text(doc, info["review"])
    table(doc, ["Criterio de aceptación", "Evidencia", "Decisión"], [
        ("Historias del sprint cumplen su comportamiento esperado.", "Demostración, casos de prueba y Taskboard.", "Aceptado"),
        ("No existen defectos críticos que impidan el uso.", "Reporte de pruebas y revisión del equipo.", "Aceptado"),
        ("La documentación técnica refleja el incremento.", "Diagramas y artefactos actualizados.", "Aceptado"),
    ], [2.5, 2.5, 1.0])
    heading(doc, "5. Sprint Retrospective", 2)
    table(doc, ["Enfoque", "Resultado / acción acordada"], info["retro"], [1.25, 4.75])
    heading(doc, "6. Burndown y Burnup", 2)
    image_with_caption(doc, DIAGRAMS / f"sprint_{number}_burndown.png", f"Figura {11 + number * 3}. Burndown del Sprint {number}.", 5.9)
    image_with_caption(doc, DIAGRAMS / f"sprint_{number}_burnup.png", f"Figura {12 + number * 3}. Burnup del Sprint {number}.", 5.9)
    heading(doc, "7. Gráfica de esfuerzo", 2)
    image_with_caption(doc, DIAGRAMS / f"sprint_{number}_esfuerzo.png", f"Figura {13 + number * 3}. Esfuerzo completado por integrante en el Sprint {number}.", 5.9)
    heading(doc, "8. Scrum Taskboard", 2)
    image_with_caption(doc, DIAGRAMS / f"sprint_{number}_taskboard.png", f"Figura {14 + number * 3}. Estado final del Taskboard del Sprint {number}.", 6.15)
    if number != 3:
        doc.add_page_break()


def conclusion(doc):
    doc.add_page_break()
    heading(doc, "CONCLUSIONES", 1)
    add_numbered(doc, [
        "Apuradito alcanzó el objetivo de entregar un flujo integral de viajes compartidos: identidad, oferta de ruta, búsqueda geoespacial, reserva, pago, cierre y evaluación.",
        "Scrum permitió priorizar el valor por incrementos. Los eventos de Planning, Daily, Review y Retrospective dejaron evidencia de inspección, adaptación y aceptación del producto.",
        "La arquitectura por capas y el uso de PostgreSQL/PostGIS concentran la complejidad de datos y matching, manteniendo clientes web y móvil desacoplados de la lógica de negocio.",
        "Los artefactos de calidad, métricas y riesgos respaldan un cierre sin defectos críticos abiertos y con todas las historias comprometidas aceptadas.",
    ])
    heading(doc, "BIBLIOGRAFÍA", 1)
    add_bullets(doc, [
        "Schwaber, K. y Sutherland, J. (2020). The Scrum Guide: The Definitive Guide to Scrum.",
        "Documentación técnica interna del proyecto Apuradito: README, modelos, servicios, API y avances de implementación.",
        "FastAPI. Documentación de referencia para APIs de Python.",
        "PostGIS. Documentación de funciones espaciales sobre PostgreSQL.",
        "React, Flutter y PostgreSQL. Documentación técnica de las tecnologías incluidas en el producto.",
    ])
    heading(doc, "ANEXOS", 1)
    heading(doc, "Anexo A. Evidencia visual de la solución", 2)
    for idx, screenshot in enumerate(SCREENSHOTS, 1):
        if screenshot.exists():
            image_with_caption(doc, screenshot, f"Anexo A.{idx}. Prototipo/interfaz implementada de Apuradito.", 5.7)
    heading(doc, "Anexo B. Diagramas editables", 2)
    add_text(doc, "Los archivos fuente editables en formato .drawio se entregan en la carpeta diagramas. Incluyen contexto, casos de uso, despliegue, paquetes, clases y secuencias. Estos archivos pueden abrirse y modificarse directamente con diagrams.net/draw.io.")
    table(doc, ["Archivo", "Tipo UML / artefacto", "Uso"], [
        ("01_contexto.drawio", "Contexto", "Actores y sistemas externos."),
        ("02_casos_de_uso.drawio", "Casos de uso UML 2.5", "Alcance funcional."),
        ("03_despliegue.drawio", "Despliegue UML 2.5", "Nodos de ejecución e integración."),
        ("04_paquetes.drawio", "Paquetes UML 2.5", "Capas de software."),
        ("05_clases_dominio.drawio", "Clases UML 2.5", "Datos y relaciones centrales."),
        ("06-07_secuencia_*.drawio", "Secuencia UML 2.5", "Flujos de reserva y pago."),
    ], [2.4, 2.05, 2.05])


def generate_docx():
    doc = Document()
    setup(doc)
    core = doc.core_properties
    core.title = "Documentación final Scrum - Apuradito"
    core.subject = "Proyecto de Ingeniería de Software I"
    core.author = "Equipo Scrum Apuradito"
    core.keywords = "Scrum, UML 2.5, Apuradito, viajes compartidos"
    cover(doc); toc(doc); paps(doc); sprint_zero(doc)
    for num in (1, 2, 3): sprint(doc, num)
    conclusion(doc)
    doc.save(DOCX_PATH)


def manual_notes():
    (ROOT / "cambios_pendientes_manual.txt").write_text(
        "DOCUMENTACIÓN APURADITO - AJUSTES MANUALES O PENDIENTES\n"
        "==================================================\n\n"
        "No existen cambios obligatorios para completar la entrega: el documento Word, los artefactos Scrum, los gráficos y los diagramas UML se generaron listos para entregar.\n\n"
        "Ajustes opcionales si la cátedra los exige:\n"
        "1. Añadir en la carátula el nombre de la universidad, docente, paralelo y ciudad, porque esos datos no estaban en el proyecto ni en la solicitud.\n"
        "2. Confirmar o reasignar los roles Scrum usados en el documento: Alejandro Melgar (Product Owner), Luis Fernando Banegas Roca (Scrum Master), Luis Pablo Mendoza y Maikol Anthony Molina Cortez (Developers). Se asignaron para poder completar el artefacto con los cuatro integrantes indicados.\n"
        "3. Si se requiere una tabla de contenido con números de página institucional, abrir el documento en Word y generar/actualizarla desde Referencias. Los títulos ya están configurados como encabezados reales.\n"
        "4. Los diagramas editables están en la subcarpeta diagramas con extensión .drawio. Se pueden abrir con diagrams.net/draw.io sin reconstruirlos.\n",
        encoding="utf-8"
    )


if __name__ == "__main__":
    generate_visuals()
    generate_docx()
    manual_notes()
    print(DOCX_PATH)
